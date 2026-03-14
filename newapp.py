"""
SRS/SDD Document Generator v2 - Offline RAG System
Features:
  - Mixed-format HRS parsing with auto-detection
  - Requirement type classification (functional, performance, interface, safety)
  - Traceability matrix (HRS → SRS → SDD)
  - Section-by-section regeneration
  - Chained SRS → SDD pipeline
  - .docx export matching user's template format

Stack: ChromaDB + SentenceTransformers + Ollama (100% offline)
"""

import os
import json
import hashlib
import re
import time
import uuid
from pathlib import Path
from datetime import datetime

from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename

import chromadb
try:
    from chromadb.api.types import EmbeddingFunction, Documents, Embeddings
    CHROMA_HAS_TYPES = True
except ImportError:
    CHROMA_HAS_TYPES = False

from sentence_transformers import SentenceTransformer
from docx import Document as DocxDocument
import fitz  # PyMuPDF
import requests


# ============================================================
# CONFIGURATION
# ============================================================

# --- LLM Backend Configuration ---
# The system auto-detects which backend is running.
# Priority: 1) Ollama  2) llama.cpp server
#
# For Ollama:     ollama serve (port 11434)
# For llama.cpp:  ./server -m model.gguf -c 8192 --port 8080 -ngl 99
#
# Override with environment variables:
#   LLM_BACKEND=ollama or LLM_BACKEND=llamacpp
#   OLLAMA_URL=http://localhost:11434
#   LLAMACPP_URL=http://localhost:8080
#   OLLAMA_MODEL=mistral:7b-instruct

OLLAMA_BASE_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")
LLAMACPP_BASE_URL = os.environ.get("LLAMACPP_URL", "http://localhost:8080")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "mistral-7b-instruct-v0.2.Q4_K_M")
LLM_BACKEND = os.environ.get("LLM_BACKEND", "auto")  # "auto", "ollama", "llamacpp"

# --- Paths ---
CHROMA_PERSIST_DIR = "./chroma_db"
UPLOAD_DIR = "./uploads"
OUTPUT_DIR = "./outputs"
PROJECTS_DIR = "./projects"

# --- Embedding Model ---
# Use "all-MiniLM-L6-v2" for HuggingFace model name
# Or use a local path like "./sentence-trans" if you downloaded manually
EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL", "all-MiniLM-L6-v2")

for d in [UPLOAD_DIR, OUTPUT_DIR, CHROMA_PERSIST_DIR, PROJECTS_DIR]:
    os.makedirs(d, exist_ok=True)


# ============================================================
# LOCAL EMBEDDING MODEL
# ============================================================
print("[*] Loading embedding model...")
embedding_model = SentenceTransformer(EMBEDDING_MODEL)

if CHROMA_HAS_TYPES:
    class LocalEmbeddingFunction(EmbeddingFunction):
        """ChromaDB-compatible embedding function (typed version)"""
        def __init__(self, model):
            self.model = model

        def __call__(self, input: Documents) -> Embeddings:
            if isinstance(input, str):
                input = [input]
            return self.model.encode(input).tolist()
else:
    class LocalEmbeddingFunction:
        """ChromaDB-compatible embedding function (fallback version)"""
        def __init__(self, model):
            self.model = model

        def __call__(self, input):
            if isinstance(input, str):
                input = [input]
            return self.model.encode(input).tolist()

embed_fn = LocalEmbeddingFunction(embedding_model)


# ============================================================
# CHROMADB COLLECTIONS
# ============================================================
print("[*] Initializing ChromaDB...")
chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)

COLLECTIONS_CONFIG = {
    "srs_templates": "Your SRS template files (structure & format reference)",
    "sdd_templates": "Your SDD template files (structure & format reference)",
    "requirements":  "HRS documents, feature specs, requirement sheets",
    "reference_docs": "Past SRS/SDD documents (writing style reference)",
}

collections = {}
for name, desc in COLLECTIONS_CONFIG.items():
    collections[name] = chroma_client.get_or_create_collection(
        name=name, embedding_function=embed_fn, metadata={"description": desc}
    )


# ============================================================
# DOCUMENT PARSERS
# ============================================================

def parse_docx(filepath):
    doc = DocxDocument(filepath)
    sections = []
    current = {"heading": "Document Start", "content": [], "level": 0}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current["content"]:
                sections.append(current)
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
            current = {"heading": text, "content": [], "level": level}
        else:
            current["content"].append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            table_text = "\n".join([" | ".join(row) for row in rows])
            current["content"].append(f"[TABLE]\n{table_text}\n[/TABLE]")

    if current["content"]:
        sections.append(current)
    return sections


def parse_pdf(filepath):
    doc = fitz.open(filepath)
    sections = []
    current_text = []
    current_heading = "Document"

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                text = " ".join(span["text"] for span in line["spans"]).strip()
                if not text:
                    continue
                font_size = max(span["size"] for span in line["spans"])
                if font_size > 13 and len(text) < 120:
                    if current_text:
                        sections.append({"heading": current_heading, "content": current_text})
                        current_text = []
                    current_heading = text
                else:
                    current_text.append(text)

    if current_text:
        sections.append({"heading": current_heading, "content": current_text})
    doc.close()
    return sections


def parse_txt(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    sections = []
    lines = content.split("\n")
    current_heading = "Document"
    current_content = []

    for line in lines:
        stripped = line.strip()
        # Detect headings: markdown style or ALL CAPS short lines
        if re.match(r'^#{1,4}\s+', stripped):
            if current_content:
                sections.append({"heading": current_heading, "content": current_content})
                current_content = []
            current_heading = re.sub(r'^#{1,4}\s+', '', stripped)
        elif stripped.isupper() and 3 < len(stripped) < 80 and not stripped.startswith(("REQ", "HR", "SR")):
            if current_content:
                sections.append({"heading": current_heading, "content": current_content})
                current_content = []
            current_heading = stripped
        elif stripped:
            current_content.append(stripped)

    if current_content:
        sections.append({"heading": current_heading, "content": current_content})
    return sections


def parse_document(filepath):
    ext = Path(filepath).suffix.lower()
    parsers = {".docx": parse_docx, ".pdf": parse_pdf, ".txt": parse_txt, ".md": parse_txt}
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(filepath)


# ============================================================
# HRS REQUIREMENT EXTRACTOR (Mixed Format Support)
# ============================================================

# Patterns for detecting requirements in mixed-format documents
REQ_PATTERNS = [
    # Numbered: REQ-001, HRS-001, SR-001, etc.
    re.compile(r'(?P<id>(?:REQ|HRS|SRS|SDD|SR|HR|FR|NFR|IR|PR|SAF)[-_]?\d{1,5})\s*[:\-–]\s*(?P<text>.+)', re.IGNORECASE),
    # "The system shall..." / "The software shall..."
    re.compile(r'(?P<text>(?:The\s+)?(?:system|software|hardware|module|component|interface|unit)\s+(?:shall|must|should|will)\s+.+)', re.IGNORECASE),
    # Numbered list: "1. ...", "1.1 ...", "a) ..."
    re.compile(r'^\s*(?P<id>\d+(?:\.\d+)*)\s*[.)]\s*(?P<text>.{20,})', re.MULTILINE),
    # Bullet with requirement language
    re.compile(r'^\s*[-•*]\s*(?P<text>(?:shall|must|should|will|provide|support|enable|ensure|maintain|implement|handle)\s+.+)', re.IGNORECASE | re.MULTILINE),
]

# Keywords for classifying requirement types
REQ_TYPE_KEYWORDS = {
    "functional": [
        "shall", "must", "process", "compute", "calculate", "generate", "display",
        "store", "retrieve", "transmit", "receive", "execute", "perform", "provide",
        "support", "enable", "allow", "handle", "manage", "control", "monitor",
        "detect", "identify", "classify", "validate", "verify", "authenticate",
        "encrypt", "decode", "encode", "convert", "transform", "filter", "route",
        "log", "record", "report", "notify", "alert", "trigger", "initiate",
    ],
    "performance": [
        "latency", "throughput", "bandwidth", "response time", "within", "millisecond",
        "microsecond", "second", "per second", "fps", "mbps", "gbps", "mhz", "ghz",
        "capacity", "maximum", "minimum", "rate", "speed", "frequency", "concurrent",
        "simultaneous", "real-time", "real time", "deadline", "utilization", "efficiency",
        "load", "scalab", "uptime", "availability", "99.9", "mtbf", "mttr",
    ],
    "interface": [
        "interface", "protocol", "api", "uart", "spi", "i2c", "pcie", "usb",
        "ethernet", "can bus", "mil-std", "arinc", "rs-232", "rs-485", "gpio",
        "hdmi", "lvds", "jtag", "tcp", "udp", "http", "mqtt", "modbus",
        "connector", "pin", "port", "bus", "link", "channel", "socket",
        "handshake", "acknowledge", "packet", "frame", "message format",
        "data format", "input", "output", "signal", "voltage", "impedance",
    ],
    "safety": [
        "safety", "fail-safe", "failsafe", "redundan", "fault", "error handling",
        "watchdog", "timeout", "recovery", "backup", "integrity", "checksum", "crc",
        "parity", "ecc", "radiation", "hardening", "critical", "hazard", "risk",
        "protection", "isolation", "guard", "limit", "boundary", "override",
        "emergency", "shutdown", "graceful degradation", "mil-std", "do-178",
        "do-254", "iec 61508", "iso 26262", "sil", "dal", "asil",
    ],
    "constraint": [
        "operating temperature", "humidity", "altitude", "vibration", "shock",
        "weight", "size", "dimension", "power consumption", "voltage range",
        "current", "watt", "ip rating", "ingress", "environmental", "emc",
        "electromagnetic", "certification", "compliance", "standard", "regulation",
        "rohs", "reach", "ce mark", "fcc", "mil-spec", "ip67", "ip68",
    ],
}


def extract_requirements(text):
    """
    Extract requirements from mixed-format text.
    Returns list of {id, text, type, confidence, source_line}
    """
    requirements = []
    seen_texts = set()
    req_counter = 1

    lines = text.split("\n")

    for line_num, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 15:
            continue

        for pattern in REQ_PATTERNS:
            match = pattern.search(line)
            if match:
                groups = match.groupdict()
                req_text = groups.get("text", line).strip()
                req_id = groups.get("id", None)

                # De-duplicate
                text_key = req_text[:80].lower()
                if text_key in seen_texts:
                    continue
                seen_texts.add(text_key)

                # Auto-assign ID if not found
                if not req_id or len(req_id) < 2:
                    req_id = f"REQ-{req_counter:04d}"
                    req_counter += 1

                # Classify type
                req_type = classify_requirement(req_text)

                requirements.append({
                    "id": req_id.upper(),
                    "text": req_text,
                    "type": req_type,
                    "source_line": line_num + 1,
                    "original": line,
                })
                break  # Stop after first pattern match per line

    return requirements


def classify_requirement(text):
    """Classify a requirement into type based on keyword matching"""
    text_lower = text.lower()
    scores = {}

    for req_type, keywords in REQ_TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        scores[req_type] = score

    if max(scores.values()) == 0:
        return "functional"  # Default

    return max(scores, key=scores.get)


def extract_requirements_from_document(filepath):
    """Parse a document and extract all requirements"""
    sections = parse_document(filepath)
    all_reqs = []

    for section in sections:
        full_text = "\n".join(section.get("content", []))
        reqs = extract_requirements(full_text)
        for req in reqs:
            req["source_section"] = section.get("heading", "Unknown")
        all_reqs.extend(reqs)

    return all_reqs


# ============================================================
# VECTOR DB OPERATIONS
# ============================================================

def ingest_document(filepath, collection_name, metadata_extra=None):
    collection = collections[collection_name]
    sections = parse_document(filepath)
    filename = Path(filepath).name
    doc_hash = hashlib.md5(filename.encode()).hexdigest()[:12]
    chunks_added = 0

    for i, section in enumerate(sections):
        content = "\n".join(section.get("content", []))
        if not content.strip():
            continue

        words = content.split()
        for j in range(0, len(words), 500):
            chunk = " ".join(words[j:j+500])
            chunk_id = f"{doc_hash}_{i}_{j}"

            meta = {
                "source_file": filename,
                "section_heading": section.get("heading", ""),
                "collection": collection_name,
                "chunk_index": j // 500,
            }
            if metadata_extra:
                meta.update(metadata_extra)

            collection.upsert(ids=[chunk_id], documents=[chunk], metadatas=[meta])
            chunks_added += 1

    return chunks_added


def retrieve_context(query, collection_names=None, n_results=5):
    if collection_names is None:
        collection_names = list(collections.keys())

    all_results = []
    for name in collection_names:
        coll = collections[name]
        try:
            count = coll.count()
            if count == 0:
                continue
            results = coll.query(
                query_texts=[query],
                n_results=min(n_results, count)
            )
            if results and results.get("documents") and results["documents"][0]:
                for i in range(len(results["documents"][0])):
                    all_results.append({
                        "text": results["documents"][0][i],
                        "metadata": results["metadatas"][0][i],
                        "distance": results["distances"][0][i],
                        "collection": name,
                    })
        except Exception as e:
            print(f"[WARNING] Error querying collection '{name}': {e}")
            continue

    all_results.sort(key=lambda x: x["distance"])
    return all_results[:n_results * 2]


# ============================================================
# LLM BACKEND (Auto-detects Ollama or llama.cpp server)
# ============================================================

def detect_backend():
    """Auto-detect which LLM server is running"""
    if LLM_BACKEND != "auto":
        return LLM_BACKEND

    # Try Ollama first
    try:
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=3)
        if resp.status_code == 200:
            return "ollama"
    except:
        pass

    # Try llama.cpp server
    try:
        resp = requests.get(f"{LLAMACPP_BASE_URL}/health", timeout=3)
        if resp.status_code == 200:
            return "llamacpp"
    except:
        pass

    # Try llama.cpp on alternative endpoints
    for port in [8080, 8081, 8000]:
        try:
            resp = requests.get(f"http://localhost:{port}/health", timeout=2)
            if resp.status_code == 200:
                return "llamacpp"
        except:
            continue

    return None


def query_llm(prompt, system_prompt="", temperature=0.3, max_retries=2):
    """Send prompt to whichever LLM backend is available"""
    backend = detect_backend()

    if backend is None:
        return ("ERROR: No LLM server found. Start one of these:\n"
                "  Ollama:     ollama serve\n"
                "  llama.cpp:  ./server -m model.gguf -c 8192 --port 8080 -ngl 99")

    for attempt in range(max_retries + 1):
        try:
            if backend == "ollama":
                return _query_ollama(prompt, system_prompt, temperature)
            else:
                return _query_llamacpp(prompt, system_prompt, temperature)
        except requests.exceptions.ConnectionError:
            if attempt == max_retries:
                return f"ERROR: Lost connection to {backend} server."
        except requests.exceptions.Timeout:
            if attempt == max_retries:
                return f"ERROR: Request timed out (10 min). Model may be too slow or prompt too large."
        except Exception as e:
            if attempt == max_retries:
                return f"ERROR: {str(e)}"
        time.sleep(2)


def _query_ollama(prompt, system_prompt, temperature):
    """Ollama API format"""
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "system": system_prompt,
        "stream": False,
        "options": {"temperature": temperature, "num_ctx": 8192, "top_p": 0.9},
    }
    resp = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json=payload, timeout=600
    )
    resp.raise_for_status()
    return resp.json()["response"]


def _query_llamacpp(prompt, system_prompt, temperature):
    """llama.cpp server API format"""
    # Build the full prompt with system prompt
    if system_prompt:
        full_prompt = f"[INST] <<SYS>>\n{system_prompt}\n<</SYS>>\n\n{prompt} [/INST]"
    else:
        full_prompt = f"[INST] {prompt} [/INST]"

    payload = {
        "prompt": full_prompt,
        "temperature": temperature,
        "n_predict": 4096,
        "top_p": 0.9,
        "stop": ["</s>", "[INST]"],
        "stream": False,
    }
    resp = requests.post(
        f"{LLAMACPP_BASE_URL}/completion",
        json=payload, timeout=600
    )
    resp.raise_for_status()
    data = resp.json()
    return data.get("content", data.get("text", ""))


def check_llm_status():
    """Check LLM server status for the UI"""
    backend = detect_backend()
    if backend == "ollama":
        return {"running": True, "backend": "ollama", "model": OLLAMA_MODEL}
    elif backend == "llamacpp":
        # Get model info from llama.cpp
        try:
            resp = requests.get(f"{LLAMACPP_BASE_URL}/props", timeout=3)
            if resp.status_code == 200:
                model_name = resp.json().get("default_generation_settings", {}).get("model", "llama.cpp model")
            else:
                model_name = "llama.cpp model"
        except:
            model_name = "llama.cpp model"
        return {"running": True, "backend": "llamacpp", "model": model_name}
    else:
        return {"running": False, "backend": None, "model": None}


# ============================================================
# SRS SECTION DEFINITIONS (IEEE 830 based)
# ============================================================

SRS_SECTIONS = [
    {"id": "1",   "title": "Introduction", "subsections": [
        {"id": "1.1", "title": "Purpose"},
        {"id": "1.2", "title": "Scope"},
        {"id": "1.3", "title": "Definitions, Acronyms, and Abbreviations"},
        {"id": "1.4", "title": "References"},
        {"id": "1.5", "title": "Overview"},
    ]},
    {"id": "2",   "title": "Overall Description", "subsections": [
        {"id": "2.1", "title": "Product Perspective"},
        {"id": "2.2", "title": "Product Functions"},
        {"id": "2.3", "title": "User Characteristics"},
        {"id": "2.4", "title": "Constraints"},
        {"id": "2.5", "title": "Assumptions and Dependencies"},
    ]},
    {"id": "3",   "title": "Specific Requirements", "subsections": [
        {"id": "3.1", "title": "Functional Requirements"},
        {"id": "3.2", "title": "Performance Requirements"},
        {"id": "3.3", "title": "Interface Requirements"},
        {"id": "3.4", "title": "Safety Requirements"},
        {"id": "3.5", "title": "Design Constraints"},
    ]},
    {"id": "4",   "title": "Traceability Matrix"},
    {"id": "5",   "title": "Appendices"},
]

SDD_SECTIONS = [
    {"id": "1",   "title": "Introduction", "subsections": [
        {"id": "1.1", "title": "Purpose"},
        {"id": "1.2", "title": "Scope"},
        {"id": "1.3", "title": "Definitions and Abbreviations"},
        {"id": "1.4", "title": "References"},
    ]},
    {"id": "2",   "title": "System Architecture", "subsections": [
        {"id": "2.1", "title": "Architecture Overview"},
        {"id": "2.2", "title": "System Context Diagram"},
        {"id": "2.3", "title": "Component Diagram"},
    ]},
    {"id": "3",   "title": "Module Design", "subsections": [
        {"id": "3.1", "title": "Module Decomposition"},
        {"id": "3.2", "title": "Module Descriptions"},
        {"id": "3.3", "title": "Module Interfaces"},
    ]},
    {"id": "4",   "title": "Data Design", "subsections": [
        {"id": "4.1", "title": "Data Structures"},
        {"id": "4.2", "title": "Data Flow"},
        {"id": "4.3", "title": "Database Design"},
    ]},
    {"id": "5",   "title": "Interface Design", "subsections": [
        {"id": "5.1", "title": "Hardware Interfaces"},
        {"id": "5.2", "title": "Software Interfaces"},
        {"id": "5.3", "title": "Communication Interfaces"},
        {"id": "5.4", "title": "User Interfaces"},
    ]},
    {"id": "6",   "title": "Detailed Design"},
    {"id": "7",   "title": "Traceability (SRS → SDD)"},
    {"id": "8",   "title": "Appendices"},
]


# ============================================================
# SYSTEM PROMPTS
# ============================================================

SRS_SYSTEM_PROMPT = """You are a senior Systems Engineer writing IEEE 830 / ISO/IEC/IEEE 29148 
compliant SRS documents for defense and aerospace projects.

RULES:
1. Follow the user's template structure EXACTLY if provided
2. Match the writing style from reference documents
3. Every requirement MUST have a unique ID (SRS-FUNC-001, SRS-PERF-001, SRS-IF-001, SRS-SAF-001)
4. Map each SRS requirement back to its source HRS requirement ID
5. Requirements must be testable, unambiguous, and atomic
6. Use "shall" for mandatory, "should" for recommended, "may" for optional
7. Include rationale where non-obvious
8. Output clean Markdown with proper heading hierarchy
9. For the traceability matrix, create a table mapping HRS-ID → SRS-ID → Verification Method"""


SDD_SYSTEM_PROMPT = """You are a senior Software Architect writing IEEE 1016 compliant 
SDD documents for defense and aerospace embedded systems.

RULES:
1. Follow the user's template structure EXACTLY if provided
2. Derive ALL design decisions from the SRS requirements
3. Every design element must trace back to an SRS requirement
4. Include component diagrams as text descriptions (UML-style)
5. Specify data structures, algorithms, and interfaces clearly
6. Document design rationale and trade-off decisions
7. Include error handling and fault tolerance design
8. For traceability: create table mapping SRS-ID → SDD Module → Design Element
9. Output clean Markdown with proper heading hierarchy"""


SECTION_REGEN_PROMPT = """You are regenerating a SPECIFIC SECTION of a {doc_type} document.
You will be given:
- The section to regenerate (ID and title)
- The full document context (other sections already generated)
- User's feedback/instructions for this section
- Relevant reference material from the knowledge base

Generate ONLY the requested section. Output clean Markdown.
Maintain consistency with the rest of the document."""


# ============================================================
# GENERATION ENGINE
# ============================================================

def build_context(project_name, requirements_text, doc_type="srs"):
    """Retrieve all relevant context from ChromaDB"""
    template_coll = "srs_templates" if doc_type == "srs" else "sdd_templates"

    template_chunks = retrieve_context(
        f"{doc_type} template document structure format sections",
        collection_names=[template_coll], n_results=8
    )
    reference_chunks = retrieve_context(
        f"{project_name} {requirements_text[:300]}",
        collection_names=["reference_docs"], n_results=5
    )
    req_chunks = retrieve_context(
        requirements_text[:500],
        collection_names=["requirements"], n_results=5
    )

    return {
        "templates": "\n\n---\n".join(
            f"[Template: {c['metadata'].get('section_heading', '')}]\n{c['text']}"
            for c in template_chunks
        ) or "No templates ingested. Use IEEE standard format.",

        "references": "\n\n---\n".join(
            f"[Ref: {c['metadata'].get('source_file', '')}]\n{c['text']}"
            for c in reference_chunks
        ) or "No reference documents available.",

        "stored_reqs": "\n\n---\n".join(
            f"[{c['metadata'].get('section_heading', '')}]\n{c['text']}"
            for c in req_chunks
        ) or "No stored requirements found.",
    }


def format_requirements_for_prompt(requirements):
    """Format extracted requirements grouped by type"""
    by_type = {}
    for req in requirements:
        rtype = req["type"]
        if rtype not in by_type:
            by_type[rtype] = []
        by_type[rtype].append(req)

    output = []
    for rtype, reqs in by_type.items():
        output.append(f"\n### {rtype.upper()} REQUIREMENTS ({len(reqs)} items)")
        for r in reqs:
            output.append(f"  [{r['id']}] {r['text']}")
            output.append(f"    Source section: {r.get('source_section', 'N/A')}")
    return "\n".join(output)


def generate_full_srs(project_name, requirements, raw_text, instructions=""):
    """Generate complete SRS document"""
    context = build_context(project_name, raw_text, "srs")
    formatted_reqs = format_requirements_for_prompt(requirements)

    prompt = f"""
PROJECT: {project_name}
DATE: {datetime.now().strftime('%Y-%m-%d')}

{'='*60}
YOUR TEMPLATE STRUCTURE (follow this format):
{'='*60}
{context['templates']}

{'='*60}
REFERENCE DOCUMENTS (match this writing style):
{'='*60}
{context['references']}

{'='*60}
STORED REQUIREMENTS FROM KNOWLEDGE BASE:
{'='*60}
{context['stored_reqs']}

{'='*60}
NEW HRS REQUIREMENTS (auto-extracted and classified):
{'='*60}
{formatted_reqs}

{'='*60}
RAW REQUIREMENTS TEXT:
{'='*60}
{raw_text[:3000]}

{'='*60}
ADDITIONAL INSTRUCTIONS:
{'='*60}
{instructions or 'Generate complete SRS with all sections.'}

IMPORTANT:
- Assign SRS IDs: SRS-FUNC-001, SRS-PERF-001, SRS-IF-001, SRS-SAF-001
- Map each SRS requirement to source HRS ID
- Include traceability matrix at the end
- Follow the template structure if available, else use IEEE 830

Generate the complete SRS document now:
"""
    return query_llm(prompt, system_prompt=SRS_SYSTEM_PROMPT, temperature=0.3)


def generate_full_sdd(project_name, srs_content, instructions=""):
    """Generate complete SDD from SRS (chained pipeline)"""
    context = build_context(project_name, srs_content[:1000], "sdd")

    prompt = f"""
PROJECT: {project_name}
DATE: {datetime.now().strftime('%Y-%m-%d')}

{'='*60}
YOUR SDD TEMPLATE STRUCTURE:
{'='*60}
{context['templates']}

{'='*60}
REFERENCE DOCUMENTS (style reference):
{'='*60}
{context['references']}

{'='*60}
SOURCE SRS DOCUMENT (design MUST satisfy these requirements):
{'='*60}
{srs_content[:6000]}

{'='*60}
ADDITIONAL INSTRUCTIONS:
{'='*60}
{instructions or 'Generate complete SDD derived from the SRS.'}

IMPORTANT:
- Every SRS requirement must have a corresponding design element
- Include SRS-ID → SDD Module traceability
- Describe architecture, modules, interfaces, data design
- Include error handling and fault tolerance design

Generate the complete SDD document now:
"""
    return query_llm(prompt, system_prompt=SDD_SYSTEM_PROMPT, temperature=0.3)


def regenerate_section(project_name, doc_type, section_id, section_title,
                       full_document, feedback, requirements_text=""):
    """Regenerate a single section based on user feedback"""
    context = build_context(project_name, section_title, doc_type)

    prompt = f"""
PROJECT: {project_name}
DOCUMENT TYPE: {doc_type.upper()}
SECTION TO REGENERATE: {section_id} - {section_title}

{'='*60}
CURRENT FULL DOCUMENT (for context):
{'='*60}
{full_document[:4000]}

{'='*60}
RELEVANT REFERENCE MATERIAL:
{'='*60}
{context['templates'][:2000]}
{context['references'][:2000]}

{'='*60}
REQUIREMENTS RELEVANT TO THIS SECTION:
{'='*60}
{requirements_text[:2000]}

{'='*60}
YOUR FEEDBACK / INSTRUCTIONS FOR THIS SECTION:
{'='*60}
{feedback}

Generate ONLY section "{section_id} - {section_title}" with improvements.
Keep the same heading level and numbering. Output Markdown.
"""
    return query_llm(prompt, system_prompt=SECTION_REGEN_PROMPT.format(doc_type=doc_type.upper()))


# ============================================================
# TRACEABILITY MATRIX GENERATOR
# ============================================================

def generate_traceability(requirements, srs_content, sdd_content=None):
    """Ask the LLM to build a traceability matrix from extracted data"""
    formatted_reqs = format_requirements_for_prompt(requirements)

    prompt = f"""
Analyze these documents and create a complete traceability matrix.

HRS REQUIREMENTS:
{formatted_reqs}

SRS DOCUMENT:
{srs_content[:4000]}

{f'SDD DOCUMENT:{chr(10)}{sdd_content[:4000]}' if sdd_content else ''}

Create a traceability matrix as a Markdown table with columns:
| HRS ID | HRS Description | SRS ID | SRS Description | {'SDD Module | ' if sdd_content else ''}Verification Method |

Rules:
- Every HRS requirement must appear in the matrix
- If an HRS has no corresponding SRS, mark as "NOT COVERED"  
- Verification methods: Test, Inspection, Analysis, Demonstration
- Be precise with ID mappings

Output ONLY the table, no other text.
"""
    return query_llm(prompt, temperature=0.1)


# ============================================================
# PROJECT STATE MANAGEMENT
# ============================================================

def save_project(project_id, data):
    path = os.path.join(PROJECTS_DIR, f"{project_id}.json")
    with open(path, "w") as f:
        json.dump(data, f, indent=2, default=str)


def load_project(project_id):
    path = os.path.join(PROJECTS_DIR, f"{project_id}.json")
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return None


def list_projects():
    projects = []
    for f in os.listdir(PROJECTS_DIR):
        if f.endswith(".json"):
            data = load_project(f.replace(".json", ""))
            if data:
                projects.append({
                    "id": data.get("id"),
                    "name": data.get("project_name"),
                    "created": data.get("created"),
                    "has_srs": bool(data.get("srs_content")),
                    "has_sdd": bool(data.get("sdd_content")),
                })
    return projects


# ============================================================
# DOCX EXPORT (basic — uses python-docx)
# ============================================================

def markdown_to_docx(markdown_text, output_path, title="Document"):
    """Convert generated Markdown to .docx with basic formatting"""
    doc = DocxDocument()

    # Title
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()

    in_table = False
    table_rows = []

    for line in markdown_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            if in_table and table_rows:
                # Flush table
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            if "---" in stripped:
                continue  # Skip separator
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue

        if in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            # Handle bold markdown
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

    # Flush any remaining table
    if table_rows:
        _add_table(doc, table_rows)

    doc.save(output_path)
    return output_path


def _add_table(doc, rows):
    """Add a table to the document"""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text

    # Bold header row
    if len(rows) > 0:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


# ============================================================
# FLASK APP
# ============================================================
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}


@app.route("/")
def index():
    return send_file("templates/index.html")


@app.route("/api/status")
def api_status():
    llm_status = check_llm_status()
    return jsonify({
        "ollama": llm_status["running"],
        "backend": llm_status["backend"],
        "model": llm_status["model"] or OLLAMA_MODEL,
        "collections": {
            name: {"count": coll.count(), "description": COLLECTIONS_CONFIG[name]}
            for name, coll in collections.items()
        }
    })


@app.route("/api/ingest", methods=["POST"])
def api_ingest():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    collection_name = request.form.get("collection", "reference_docs")

    if collection_name not in collections:
        return jsonify({"error": f"Invalid collection"}), 400

    filename = secure_filename(file.filename)
    if Path(filename).suffix.lower() not in ALLOWED_EXTENSIONS:
        return jsonify({"error": "Unsupported file type"}), 400

    filepath = os.path.join(UPLOAD_DIR, filename)
    file.save(filepath)

    try:
        chunks = ingest_document(filepath, collection_name)
        return jsonify({
            "status": "success", "filename": filename,
            "collection": collection_name, "chunks_created": chunks
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/extract-requirements", methods=["POST"])
def api_extract_requirements():
    """Upload HRS file and extract requirements"""
    if "file" in request.files:
        file = request.files["file"]
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_DIR, filename)
        file.save(filepath)
        reqs = extract_requirements_from_document(filepath)
        raw_text = "\n".join(
            "\n".join(s.get("content", []))
            for s in parse_document(filepath)
        )
    elif request.json and "text" in request.json:
        raw_text = request.json["text"]
        reqs = extract_requirements(raw_text)
    else:
        return jsonify({"error": "Provide file or text"}), 400

    # Summary
    type_counts = {}
    for r in reqs:
        type_counts[r["type"]] = type_counts.get(r["type"], 0) + 1

    return jsonify({
        "status": "success",
        "total": len(reqs),
        "by_type": type_counts,
        "requirements": reqs,
        "raw_text": raw_text[:5000],
    })


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """Generate SRS or SDD"""
    try:
        data = request.json
        doc_type = data.get("doc_type", "srs")
        project_name = data.get("project_name", "Untitled")
        requirements = data.get("requirements", [])
        raw_text = data.get("raw_requirements_text", "")
        instructions = data.get("instructions", "")
        project_id = data.get("project_id") or str(uuid.uuid4())[:8]

        if doc_type == "srs":
            content = generate_full_srs(project_name, requirements, raw_text, instructions)
        elif doc_type == "sdd":
            project_data = load_project(project_id)
            srs_content = data.get("srs_content", "")
            if not srs_content and project_data:
                srs_content = project_data.get("srs_content", "")
            if not srs_content:
                return jsonify({"error": "SDD generation requires SRS content. Generate SRS first."}), 400
            content = generate_full_sdd(project_name, srs_content, instructions)
        else:
            return jsonify({"error": "doc_type must be 'srs' or 'sdd'"}), 400

        if content.startswith("ERROR:"):
            return jsonify({"error": content}), 500

        # Save .docx
        docx_filename = f"{project_name.replace(' ', '_')}_{doc_type.upper()}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        markdown_to_docx(content, docx_path, f"{project_name} — {doc_type.upper()}")

        # Save project state
        project_data = load_project(project_id) or {
            "id": project_id,
            "project_name": project_name,
            "created": datetime.now().isoformat(),
            "requirements": requirements,
            "raw_requirements": raw_text[:10000],
        }
        if doc_type == "srs":
            project_data["srs_content"] = content
            project_data["srs_docx"] = docx_filename
        else:
            project_data["sdd_content"] = content
            project_data["sdd_docx"] = docx_filename
        save_project(project_id, project_data)

        return jsonify({
            "status": "success",
            "project_id": project_id,
            "doc_type": doc_type,
            "content": content,
            "docx_filename": docx_filename,
        })
    except Exception as e:
        print(f"[ERROR] Generate failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-sdd-from-srs", methods=["POST"])
def api_generate_sdd_pipeline():
    """Chained SRS → SDD pipeline"""
    data = request.json
    project_id = data.get("project_id")
    instructions = data.get("instructions", "")

    project_data = load_project(project_id)
    if not project_data or not project_data.get("srs_content"):
        return jsonify({"error": "No SRS found for this project. Generate SRS first."}), 400

    sdd_content = generate_full_sdd(
        project_data["project_name"],
        project_data["srs_content"],
        instructions
    )

    docx_filename = f"{project_data['project_name'].replace(' ', '_')}_SDD.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
    markdown_to_docx(sdd_content, docx_path, f"{project_data['project_name']} — SDD")

    project_data["sdd_content"] = sdd_content
    project_data["sdd_docx"] = docx_filename
    save_project(project_id, project_data)

    return jsonify({
        "status": "success",
        "project_id": project_id,
        "content": sdd_content,
        "docx_filename": docx_filename,
    })


@app.route("/api/regenerate-section", methods=["POST"])
def api_regenerate_section():
    """Regenerate a specific section"""
    data = request.json
    project_id = data.get("project_id")
    doc_type = data.get("doc_type", "srs")
    section_id = data.get("section_id")
    section_title = data.get("section_title")
    feedback = data.get("feedback", "")

    project_data = load_project(project_id)
    if not project_data:
        return jsonify({"error": "Project not found"}), 404

    full_doc = project_data.get(f"{doc_type}_content", "")
    raw_reqs = project_data.get("raw_requirements", "")

    new_section = regenerate_section(
        project_data["project_name"], doc_type,
        section_id, section_title,
        full_doc, feedback, raw_reqs
    )

    return jsonify({
        "status": "success",
        "section_id": section_id,
        "content": new_section,
    })


@app.route("/api/traceability", methods=["POST"])
def api_traceability():
    """Generate traceability matrix"""
    data = request.json
    project_id = data.get("project_id")

    project_data = load_project(project_id)
    if not project_data:
        return jsonify({"error": "Project not found"}), 404

    requirements = project_data.get("requirements", [])
    srs_content = project_data.get("srs_content", "")
    sdd_content = project_data.get("sdd_content")

    matrix = generate_traceability(requirements, srs_content, sdd_content)

    return jsonify({"status": "success", "matrix": matrix})


@app.route("/api/download/<filename>")
def api_download(filename):
    filepath = os.path.join(OUTPUT_DIR, secure_filename(filename))
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return jsonify({"error": "File not found"}), 404


@app.route("/api/projects")
def api_projects():
    return jsonify(list_projects())


@app.route("/api/project/<project_id>")
def api_project_detail(project_id):
    data = load_project(project_id)
    if data:
        return jsonify(data)
    return jsonify({"error": "Not found"}), 404


@app.route("/api/search", methods=["POST"])
def api_search():
    data = request.json
    results = retrieve_context(
        data.get("query", ""),
        collection_names=data.get("collections"),
        n_results=data.get("n_results", 5)
    )
    return jsonify({"results": results})


@app.route("/api/sections/<doc_type>")
def api_sections(doc_type):
    if doc_type == "srs":
        return jsonify(SRS_SECTIONS)
    elif doc_type == "sdd":
        return jsonify(SDD_SECTIONS)
    return jsonify({"error": "Invalid type"}), 400


if __name__ == "__main__":
    llm_info = check_llm_status()
    print("\n" + "=" * 60)
    print("  SRS/SDD RAG Generator v2")
    print(f"  LLM Backend: {llm_info['backend'] or 'NOT DETECTED'}")
    print(f"  LLM Model: {llm_info['model'] or 'N/A'}")
    print(f"  LLM Status: {'CONNECTED' if llm_info['running'] else 'NOT RUNNING'}")
    if not llm_info['running']:
        print("")
        print("  ⚠ Start your LLM server:")
        print(f"    Ollama:     ollama serve")
        print(f"    llama.cpp:  ./server -m model.gguf -c 8192 --port 8080 -ngl 99")
    print(f"  ChromaDB: {CHROMA_PERSIST_DIR}")
    print(f"  Embedding: {EMBEDDING_MODEL}")
    print(f"  Web UI: http://localhost:5000")
    print("=" * 60 + "\n")
    app.run(host="0.0.0.0", port=5000, debug=True)
